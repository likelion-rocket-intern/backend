from sqlmodel import Session
from konlpy.tag import *
from app.crud.resume import resume_crud
from app.models.resume import Resume, ResumeKeyword, ResumeEmbedding
from typing import Optional, List
from fastapi import HTTPException
from app.core.config import settings
from app.models.embedding import Embedding
from app.utils.word2vec import WordVectorModelLoader
import json
from langchain_openai import OpenAIEmbeddings
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from docx import Document as DocxDocument
from langchain_community.document_loaders import PyMuPDFLoader, Docx2txtLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from app.utils.storage import delete_resume
from app.repository.sql_embedding_repository import SqlEmbeddingRepository
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
import pandas as pd


class ResumeService:
    def __init__(self):
        # HuggingFace Embedding 모델 벡터 차원수가 1024 이므로 맞춰주어야합니다.
        # 만약 HuggingFace 모델을 사용하신다면 변경하시면 됩니다.
        """
        from langchain_huggingface.embeddings import HuggingFaceEmbeddings

        self.embeddings = HuggingFaceEmbeddings(
                model_name = settings.EMBBEDING_MODEL,
            )
        """
        
        self.embeddings = OpenAIEmbeddings(
            model=settings.OPENAI_EMBEDDING_MODEL,
            api_key=settings.OPENAI_API_KEY,
            dimensions=1024
        )

    def create(
        self,
        db: Session,
        *,
        resume: Resume
    ) -> Resume:
        return resume_crud.create(db=db, resume=resume)

    def get_by_id(
        self,
        db: Session,
        resume_id: int,
        user_id: int
    ) -> Optional[Resume]:
        resume = resume_crud.get_by_id(db=db, resume_id=resume_id)
        if resume.user_id != user_id:
            raise HTTPException(
                status_code=403,
                detail="해당 이력서는 접근 권한이 없습니다."
            )
        if not resume:
            raise HTTPException(
                status_code=404, 
                detail="Resume not found")
        return resume

    def get_by_user_id(
        self,
        db: Session,
        user_id: int
    ) -> List[Resume]:
        resumes = resume_crud.get_by_user_id(db=db, user_id=user_id)
        return resumes

    def get_with_user(
        self,
        db: Session,
        resume_id: int
    ) -> Optional[Resume]:
        resume = resume_crud.get_with_user(db=db, resume_id=resume_id)
        if not resume:
            raise HTTPException(status_code=404, detail="Resume not found")
        return resume

    def delete(
        self,
        db: Session,
        resume_id: int
    ) -> None:
        resume = self.get_by_id(db=db, resume_id=resume_id)
        
        # 스토리지에서 파일 삭제
        delete_resume(resume.file_path)
        
        # 이력서 삭제 (cascade로 인해 관련 임베딩도 자동 삭제됨)
        if resume:
            db.delete(resume)
            db.commit()

    # def parse_resume(self, file_path: str, filename: str) -> List[Document]:
    #     file_extenstion = filename.split(".")[-1].lower()

    #     if file_extenstion == "pdf":
    #         loader = PyPDFLoader(file_path)
    #         return loader.load()
    #     elif file_extenstion in ["doc", "docx"]:
    #         doc = DocxDocument(file_path)
    #         full_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    #         return [Document(page_content=full_text)]
    #     else:
    #         raise HTTPException(status_code=400, detail="Unsupported file type")

    # def process_resume_file(self, file_path: str, filename: str) -> None:
    #     document = self.parse_resume(file_path, filename)

    #     full_text = ""
    #     for doc in document:
    #         full_text += doc.page_content + "\n"

    #     return full_text
    
    def parse_and_chunk_resume(self, file_path: str, filename: str, original_filename: str) -> List[Document]:
        file_extenstion = original_filename.split(".")[-1].lower()
        
        # 텍스트 분할기 설정
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,  # 크기를 늘려서 의미단위 보존
            chunk_overlap=100,  # 겹침을 줄임
            separators=[
                "\n**",          # 섹션 제목 (볼드)
                "\n✓",           # 프로젝트 내 항목들
                "\n\n",          # 단락 구분
                "\n",            # 줄바꿈
                ".",             # 마침표
                " "              # 공백
            ]
        )
            
        if file_extenstion == 'pdf':
            loader = PyMuPDFLoader(file_path)
        elif file_extenstion in ['doc', 'docx']:
            loader = Docx2txtLoader(file_path)

        similar_words_list = self.analysis_keywords(loader.load())
        
        # 로드와 동시에 청킹
        documents = loader.load_and_split(text_splitter=text_splitter)
        # print(f"총 청크 수: {len(documents)}")
        # print("\n" + "="*50)

        # for i, doc in enumerate(documents):
        #     print(f"청크 {i+1}:")
        #     print(f"길이: {len(doc.page_content)}자")
        #     print(f"내용: {doc.page_content[:200]}...")
        #     print("-" * 30)
        
        return documents, similar_words_list
        
    def create_embeddings(self, chunks: List[Document]) -> List[List[float]]:
        try:
            texts = [chunk.page_content for chunk in chunks]

            vectors = self.embeddings.embed_documents(texts)
            return vectors
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error creating embeddings: {str(e)}")
        
    def save_embeddings_to_db(
        self, 
        db: Session, 
        resume_id: int, 
        chunks: List[Document], 
        vectors: List[List[float]]
    ) -> List[ResumeEmbedding]:
        """
        청크와 임베딩을 DB에 저장
        """
        return resume_crud.bulk_create_embeddings(
            db=db,
            resume_id=resume_id,
            chunks=chunks,
            vectors=vectors
        )
        
    # def process_resume_with_embeddings(self, file_path: str, filename: str) -> tuple[List[Document], List[List[float]]]:
    #     # 1. 파싱 & 청킹
    #     chunks = self.parse_and_chunk_resume(file_path, filename)
        
    #     # 2. 임베딩 생성
    #     vectors = self.create_embeddings(chunks)

    #     print(f"청크 수: {len(chunks)}")
    #     print(f"벡터 수: {len(vectors)}")
    #     print(f"벡터 차원: {len(vectors[0])}")
        
    #     return chunks, vectors
    
    # def process_and_save_resume(
    #     self, 
    #     db: Session, 
    #     file_path: str, 
    #     filename: str, 
    #     resume_id: int = 1  # 테스트용 고정
    # ):
    #     """
    #     이력서 처리하고 DB에 저장하는 전체 플로우
    #     """
    #     # 1. 파싱 & 청킹 & 임베딩
    #     chunks, vectors = self.process_resume_with_embeddings(file_path, filename)
        
    #     # 2. DB에 저장
    #     embedding_records = self.save_embeddings_to_db(db, resume_id, chunks, vectors)
        
    #     print(f"✅ {len(embedding_records)}개의 임베딩이 DB에 저장되었습니다!")
        
    #     return {
    #         "resume_id": resume_id,
    #         "chunks_count": len(chunks),
    #         "embeddings_saved": len(embedding_records)
    #     }

    def _get_itemwise_scores(
            self, resume_vectors: list, dataset_embeddings: list[Embedding], key_name: str
    ) -> list[dict]:
        """
        이력서 벡터와 데이터셋을 비교하여, 각 항목별 점수를 계산하고 정렬합니다.
        """
        if not dataset_embeddings or not resume_vectors:
            return []

        # 1. 데이터셋에서 벡터와 이름(또는 속성)을 분리합니다.
        target_vectors = np.array([e.embedding for e in dataset_embeddings])
        target_data_list = [e.extra_data for e in dataset_embeddings]
        
        target_identifiers = [data.get(key_name, "") for data in target_data_list]

        # 2. 코사인 유사도 계산
        similarity_matrix = cosine_similarity(np.array(resume_vectors), target_vectors)

        # 3. 각 데이터셋 항목별 '최대' 유사도 점수를 추출합니다.
        # axis=0은 각 열(항목)에서 최대값을 찾는다는 의미입니다.
        max_similarities = np.max(similarity_matrix, axis=0)
        max_similarities_positive = np.maximum(0, max_similarities)

        # 결과를 pandas DataFrame으로 변환해서 계산된 모든 결과를 (이름, 점수) 형태로 DF에 담습니다.
        df = pd.DataFrame({
            'identifier': target_identifiers,
            'score': max_similarities_positive,
            'data': target_data_list
        })

        # 정규분포 임계치 설정 (qualified_df)
        mean_score = df['score'].mean()
        std_score = df['score'].std()
        std_multiplier = 0.8
        threshold = mean_score + std_multiplier * std_score

        qualified_df = df[df['score'] >= threshold]

        if qualified_df.empty:
            return []

        # 동일한 이름으로 그룹화하여 최고점수를 추출합니다. (가장 높게 나온 유사도 점수를 가져와서 전문성 파악할 시 유리.. 다른 방식도 고려 필요 )
        grouped_scores = qualified_df.groupby('identifier').agg(
            score=('score', 'max'),
            data=('data', 'first')
        ).reset_index()

        # 4. 점수들을 백분위로 정규화합니다 (총합이 100이 되도록).
        total_similarity = grouped_scores['score'].sum()
        if total_similarity > 0:
            grouped_scores['score'] = (grouped_scores['score'] / total_similarity) * 100
        else:
            grouped_scores['score'] = 0
        
        # 5. 점수 순으로 정렬하고 결과를 반환합니다.
        sorted_results = grouped_scores.sort_values(by='score', ascending=False)
        final_results = []
        for _, row in sorted_results.iterrows():
            result_item = row['data']
            result_item['score'] = round(row['score'], 2)
            final_results.append(result_item)
            
        return final_results

    def analyze_resume_fitness(self, db: Session, resume_vectors: list) -> dict:
        """
        이력서 벡터를 '직무' 및 '이력서 평가' 데이터셋과 비교하여 종합 분석 결과를 반환합니다.
        """
        embedding_repo = SqlEmbeddingRepository(db)

        # 1. 직무 적합성 분석 (job 타입 데이터셋과 비교)
        job_embeddings = embedding_repo.get_all_by_type("job")
        job_fitness_scores = self._get_itemwise_scores(resume_vectors, job_embeddings, key_name='name')

        # 2. 이력서 강점/보완점 분석 (resume 타입 데이터셋과 비교)
        # 2-1 강점 분석
        strengths_embeddings = embedding_repo.get_all_by_type("strengths")
        strengths_evaluation_scores = self._get_itemwise_scores(resume_vectors, strengths_embeddings, key_name='attribute')

        # 2-2 보완점 분석
        weaknesses_embeddings = embedding_repo.get_all_by_type("weaknesses")
        weaknesses_evaluation_scores = self._get_itemwise_scores(resume_vectors, weaknesses_embeddings, key_name='attribute')

        # 3. 최종 결과를 구조화하여 반환합니다.
        return {
            "job_fitness": job_fitness_scores,
            "resume_evaluation": {
                    "strengths": strengths_evaluation_scores,
                    "weaknesses": weaknesses_evaluation_scores
                }
            }

    def analysis_keywords(self, documents: Document):

        json_keywords_list = WordVectorModelLoader().get_json_keywords_list()
        model = WordVectorModelLoader().get_model()
        SIMILARITY_THRESHOLD = settings.SIMILARITY_THRESHOLD

        full_text = ""
        for doc in documents:
            full_text += doc.page_content + "\n"

        okt = Okt()
        okt_morphs = okt.morphs(full_text)

        resume_dict = {}
        for word in okt_morphs:
            resume_dict[word] = resume_dict.get(word, 0) + 1

        similar_words_list = []

        try:
            for resume_word, freq in resume_dict.items():
                for criterion_keyword in json_keywords_list:
                    try:
                        similarity = model.similarity(resume_word, criterion_keyword)
                        if similarity >= SIMILARITY_THRESHOLD:
                            similar_words_list.append({
                                'keyword': resume_word,
                                'similar_to': criterion_keyword,
                                'similarity': similarity,
                                'frequency': freq
                            })
                    except KeyError:
                        continue
            sorted_list = sorted(similar_words_list, 
                                key=lambda x: (round(x['similarity'], 4), x['frequency']), 
                                reverse=True)

            return sorted_list

        except Exception as e:
            print(f"오류 발생: {e}")

resume_service = ResumeService()
